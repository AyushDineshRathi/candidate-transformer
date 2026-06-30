import os
import pytest
from fpdf import FPDF
from docx import Document
from src.extractors.resume_extractor import extract_from_resume

@pytest.fixture(scope="module")
def resume_fixtures(tmp_path_factory):
    d = tmp_path_factory.mktemp("resumes")
    
    # 1. Create PDF fixture
    pdf_path = str(d / "test_resume.pdf")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # The first short non-empty line without digits/emails/phones will be full_name
    pdf.cell(200, 10, txt="John Doe", ln=1)
    pdf.cell(200, 10, txt="johndoe@example.com", ln=1)
    pdf.cell(200, 10, txt="+1 (503) 555-0142", ln=1)
    pdf.cell(200, 10, txt="", ln=1)
    pdf.cell(200, 10, txt="SKILLS", ln=1)
    pdf.cell(200, 10, txt="Python, SQL, C++", ln=1)
    pdf.cell(200, 10, txt="", ln=1)
    pdf.cell(200, 10, txt="Education", ln=1)
    pdf.cell(200, 10, txt="University of Oregon", ln=1)
    pdf.output(pdf_path)
    
    # 2. Create DOCX fixture
    docx_path = str(d / "test_resume.docx")
    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("jane.smith@email.com")
    doc.add_paragraph("+1-503-555-0199")
    doc.add_paragraph("Technical Skills")
    doc.add_paragraph("Java • Go • React")
    doc.add_paragraph("Education")
    doc.add_paragraph("Oregon State University")
    doc.save(docx_path)
    
    # 3. Create empty PDF fixture
    empty_pdf_path = str(d / "empty_resume.pdf")
    pdf_empty = FPDF()
    pdf_empty.add_page()
    pdf_empty.output(empty_pdf_path)
    
    # 4. Create fixture mirroring the bug
    bug_pdf_path = str(d / "bug_resume.pdf")
    pdf_bug = FPDF()
    pdf_bug.add_page()
    pdf_bug.set_font("Arial", size=12)
    pdf_bug.cell(200, 10, txt="Ayush Rathi", ln=1)
    pdf_bug.cell(200, 10, txt="+91 7249381902", ln=1)
    pdf_bug.cell(200, 10, txt="", ln=1)
    pdf_bug.cell(200, 10, txt="Technical Skills", ln=1)
    pdf_bug.cell(200, 10, txt="Languages: Python, Java", ln=1)
    pdf_bug.cell(200, 10, txt="Frameworks: PyTorch, TensorFlow", ln=1)
    pdf_bug.cell(200, 10, txt="", ln=1)
    pdf_bug.cell(200, 10, txt="- Achievements:", ln=1)
    pdf_bug.cell(200, 10, txt="Solved 600+ Problems On Leetcode", ln=1)
    pdf_bug.cell(200, 10, txt="3rd Place At Intellectus", ln=1)
    pdf_bug.cell(200, 10, txt="", ln=1)
    pdf_bug.cell(200, 10, txt="Certifications", ln=1)
    pdf_bug.cell(200, 10, txt="OCI 2025 - Generative AI Professional", ln=1)
    pdf_bug.output(bug_pdf_path)
    
    return {
        "pdf": pdf_path,
        "docx": docx_path,
        "empty_pdf": empty_pdf_path,
        "bug_pdf": bug_pdf_path
    }

def test_extract_from_pdf(resume_fixtures):
    ext = extract_from_resume(resume_fixtures["pdf"])
    assert ext is not None
    assert ext.candidate_id == "RESUME-test_resume.pdf"
    
    # Full Name
    assert ext.full_name is not None
    assert ext.full_name[0] == "John Doe"
    assert ext.full_name[1].confidence == 0.6
    
    # Emails
    emails = [e[0] for e in ext.emails]
    assert "johndoe@example.com" in emails
    
    # Phones
    phones = [p[0] for p in ext.phones]
    assert "+15035550142" in phones
    
    # Skills
    skills = [s[0].lower() for s in ext.skills]
    assert "python" in skills
    assert "sql" in skills
    assert "c++" in skills
    
    # Education
    edu = [e[0]["institution"] for e in ext.education]
    assert "University of Oregon" in edu

def test_extract_from_docx(resume_fixtures):
    ext = extract_from_resume(resume_fixtures["docx"])
    assert ext is not None
    
    assert ext.full_name[0] == "Jane Smith"
    
    emails = [e[0] for e in ext.emails]
    assert "jane.smith@email.com" in emails
    
    phones = [p[0] for p in ext.phones]
    assert "+15035550199" in phones
    
    skills = [s[0].lower() for s in ext.skills]
    assert "java" in skills
    assert "go" in skills
    assert "react" in skills
    
    edu = [e[0]["institution"] for e in ext.education]
    assert "Oregon State University" in edu

def test_extract_from_empty_pdf(resume_fixtures):
    ext = extract_from_resume(resume_fixtures["empty_pdf"])
    assert ext is not None
    assert ext.full_name is None
    assert len(ext.emails) == 0
    assert len(ext.phones) == 0
    assert len(ext.skills) == 0

def test_extract_from_corrupted_file(tmp_path):
    bad_file = tmp_path / "bad.pdf"
    bad_file.write_text("not a real pdf file")
    ext = extract_from_resume(str(bad_file))
    assert ext is None

def test_extract_bug_resume_sections(resume_fixtures):
    ext = extract_from_resume(resume_fixtures["bug_pdf"])
    assert ext is not None
    
    skills = [s[0].lower() for s in ext.skills]
    assert "python" in skills
    assert "java" in skills
    assert "pytorch" in skills
    assert "tensorflow" in skills
    
    assert "languages: python" not in skills
    assert "languages" not in skills
    
    assert "solved 600+ problems on leetcode" not in skills
    assert "3rd place at intellectus" not in skills
    assert "oci 2025 - generative ai professional" not in skills
